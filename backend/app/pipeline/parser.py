"""File (PDF/DOCX) -> raw text. Keep this dumb; the LLM does structuring. No OCR."""

import re
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber
from docx import Document

_MIN_TEXT_LENGTH = 200
_MAX_TEXT_LENGTH = 15_000
_HEAD_LENGTH = 9_000
_TAIL_LENGTH = 5_000


class ParseError(Exception):
    """Raised when text can't be extracted from a file — surfaces as a 4xx to the caller."""


def _extract_pdf_pdfplumber(path: str) -> str:
    with pdfplumber.open(path) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def _extract_pdf_pymupdf(path: str) -> str:
    with fitz.open(path) as doc:
        return "\n".join(page.get_text() for page in doc)


def _extract_docx(path: str) -> str:
    document = Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\x00", "")).strip()


def _truncate(text: str) -> str:
    if len(text) <= _MAX_TEXT_LENGTH:
        return text
    return text[:_HEAD_LENGTH] + "\n[...truncated...]\n" + text[-_TAIL_LENGTH:]


def extract_text(path: str) -> str:
    suffix = Path(path).suffix.lower()

    if suffix == ".pdf":
        text = _extract_pdf_pdfplumber(path)
        if len(text) < _MIN_TEXT_LENGTH:
            text = _extract_pdf_pymupdf(path)
        if len(text) < _MIN_TEXT_LENGTH:
            raise ParseError("no text layer — scanned PDF?")
    elif suffix == ".docx":
        text = _extract_docx(path)
    else:
        raise ParseError(f"unsupported file type: {suffix or '(none)'}")

    return _truncate(_normalize(text))
